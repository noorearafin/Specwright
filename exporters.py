"""Exporters for test cases. Six formats, all write to the project root.

- CSV        → test_cases.csv          (flat, universal)
- Excel      → test_cases.xlsx         (formatted, with summary sheet)
- Jira CSV   → test_cases.jira.csv     (Jira + Xray importable)
- TestRail   → test_cases.testrail.csv (TestRail CSV importer)
- HTML       → test_cases.html         (self-contained, styled)
- Markdown   → test_cases.md           (summary table + per-case detail)
"""

from __future__ import annotations

import csv
import html
from collections import Counter
from pathlib import Path


def run_exports(cases: list[dict], out_dir: Path, formats: list[str]) -> None:
    """Dispatch test cases to each requested export format.

    Called after Stage 2 and again on every inline edit in the Streamlit UI.
    Unknown format names are warned and skipped so a typo doesn't abort everything.
    """
    # Map format names (from config / UI) to their export functions
    handlers = {
        "csv": export_csv,
        "excel": export_excel,
        "jira": export_jira,
        "testrail": export_testrail,
        "html": export_html,
        "markdown": export_markdown,
    }
    print("▶ Exporting test cases...")
    for fmt in formats:
        fn = handlers.get(fmt)
        if not fn:
            print(f"  ⚠ unknown format: {fmt}")
            continue
        path = fn(cases, out_dir)
        print(f"  ✓ {fmt:<10} → {path.name}")
    print()


# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────
def _fmt_steps(steps: list[dict]) -> str:
    """Format steps as '1. Action (data: X)\\n2. ...'"""
    out = []
    for i, s in enumerate(steps, 1):
        line = f"{i}. {s.get('action', '')}"
        if s.get("data"):
            line += f" [data: {s['data']}]"
        out.append(line)
    return "\n".join(out)


def _fmt_preconditions(pcs: list[str]) -> str:
    return "\n".join(f"- {p}" for p in pcs) if pcs else ""


def _priority_to_jira(p: str) -> str:
    return {"P0": "Highest", "P1": "High", "P2": "Medium"}.get(p, "Medium")


# ──────────────────────────────────────────────────────────────────────────────
# 1. Plain CSV
# ──────────────────────────────────────────────────────────────────────────────
def export_csv(cases: list[dict], out_dir: Path) -> Path:
    path = out_dir / "test_cases.csv"
    with path.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow([
            "ID", "Requirement", "Title", "Type", "Target", "Priority", "Page",
            "Preconditions", "Steps", "Expected", "Automatable",
        ])
        for c in cases:
            w.writerow([
                c.get("id", ""), c.get("requirement_id", ""), c.get("title", ""),
                c.get("type", ""), c.get("target", ""), c.get("priority", ""),
                c.get("page", ""),
                _fmt_preconditions(c.get("preconditions", [])),
                _fmt_steps(c.get("steps", [])),
                c.get("expected", ""),
                "Yes" if c.get("automatable", True) else "No",
            ])
    return path


# ──────────────────────────────────────────────────────────────────────────────
# 2. Excel (.xlsx) — cases sheet + summary sheet
# ──────────────────────────────────────────────────────────────────────────────
def export_excel(cases: list[dict], out_dir: Path) -> Path:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise SystemExit("Excel export needs openpyxl: pip install openpyxl")

    path = out_dir / "test_cases.xlsx"
    wb = Workbook()

    # Sheet 1: Cases
    ws = wb.active
    ws.title = "Test Cases"
    headers = ["ID", "Requirement", "Title", "Type", "Target", "Priority", "Page",
               "Preconditions", "Steps", "Expected", "Automatable"]
    ws.append(headers)

    header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(vertical="center")

    priority_colors = {"P0": "E74C3C", "P1": "F39C12", "P2": "95A5A6"}

    for c in cases:
        ws.append([
            c.get("id", ""), c.get("requirement_id", ""), c.get("title", ""),
            c.get("type", ""), c.get("target", ""), c.get("priority", ""),
            c.get("page", ""),
            _fmt_preconditions(c.get("preconditions", [])),
            _fmt_steps(c.get("steps", [])),
            c.get("expected", ""),
            "Yes" if c.get("automatable", True) else "No",
        ])
        # Color priority cell
        prio = c.get("priority", "")
        if prio in priority_colors:
            cell = ws.cell(row=ws.max_row, column=6)
            cell.fill = PatternFill(start_color=priority_colors[prio],
                                     end_color=priority_colors[prio], fill_type="solid")
            cell.font = Font(color="FFFFFF", bold=True)
            cell.alignment = Alignment(horizontal="center")

    # Column widths
    widths = [10, 12, 45, 14, 10, 10, 18, 30, 50, 40, 12]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # Wrap multi-line cells
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    ws.freeze_panes = "A2"

    # Sheet 2: Summary
    ws2 = wb.create_sheet("Summary")
    ws2.append(["Metric", "Count"])
    for cell in ws2[1]:
        cell.fill = header_fill
        cell.font = header_font

    ws2.append(["Total cases", len(cases)])
    ws2.append(["Automatable", sum(1 for c in cases if c.get("automatable", True))])
    ws2.append([])
    ws2.append(["By Priority", ""])
    for p, n in Counter(c.get("priority", "?") for c in cases).most_common():
        ws2.append([f"  {p}", n])
    ws2.append([])
    ws2.append(["By Type", ""])
    for t, n in Counter(c.get("type", "?") for c in cases).most_common():
        ws2.append([f"  {t}", n])
    ws2.append([])
    ws2.append(["By Target", ""])
    for tg, n in Counter(c.get("target", "?") for c in cases).most_common():
        ws2.append([f"  {tg}", n])

    ws2.column_dimensions["A"].width = 24
    ws2.column_dimensions["B"].width = 12

    wb.save(path)
    return path


# ──────────────────────────────────────────────────────────────────────────────
# 3. Jira CSV (Xray-compatible Test issue format)
# ──────────────────────────────────────────────────────────────────────────────
def export_jira(cases: list[dict], out_dir: Path) -> Path:
    path = out_dir / "test_cases.jira.csv"
    with path.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow([
            "Issue Type", "Summary", "Priority", "Description", "Labels",
            "Test Type", "Manual Test Steps",
        ])
        for c in cases:
            desc = _build_jira_description(c)
            labels = " ".join(filter(None, [
                c.get("type", ""),
                c.get("target", ""),
                c.get("priority", ""),
                c.get("requirement_id", ""),
            ]))
            steps_combined = _fmt_steps(c.get("steps", [])) + f"\n\nExpected: {c.get('expected', '')}"
            w.writerow([
                "Test",
                f"[{c.get('id', '')}] {c.get('title', '')}",
                _priority_to_jira(c.get("priority", "P2")),
                desc,
                labels,
                "Manual" if not c.get("automatable", True) else "Automated",
                steps_combined,
            ])
    return path


def _build_jira_description(c: dict) -> str:
    """Format a test case as Jira wiki markup for the Description field.

    Uses Jira's *bold* and # ordered-list syntax so the imported ticket
    renders nicely in the Jira UI. This is the format Xray expects when
    importing Test issues via CSV.
    """
    lines = [
        f"*Requirement:* {c.get('requirement_id', 'n/a')}",
        f"*Type:* {c.get('type', 'n/a')}  |  *Target:* {c.get('target', 'n/a')}",
        "",
    ]
    pcs = c.get("preconditions") or []
    if pcs:
        lines.append("*Preconditions:*")
        lines += [f"* {p}" for p in pcs]
        lines.append("")
    lines.append("*Steps:*")
    for i, s in enumerate(c.get("steps", []), 1):
        # Jira wiki uses # for numbered list items
        line = f"# {s.get('action', '')}"
        if s.get("data"):
            line += f" (data: {s['data']})"
        lines.append(line)
    lines.append("")
    lines.append(f"*Expected:* {c.get('expected', '')}")
    return "\n".join(lines)


# ──────────────────────────────────────────────────────────────────────────────
# 4. TestRail CSV (standard TestRail importer columns)
# ──────────────────────────────────────────────────────────────────────────────
def export_testrail(cases: list[dict], out_dir: Path) -> Path:
    path = out_dir / "test_cases.testrail.csv"
    tr_priority = {"P0": "Critical", "P1": "High", "P2": "Medium"}
    with path.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow([
            "Section", "Title", "Template", "Type", "Priority",
            "Preconditions", "Steps", "Expected Result", "References",
        ])
        for c in cases:
            w.writerow([
                c.get("requirement_id", "General"),
                c.get("title", ""),
                "Test Case (Text)",
                c.get("type", "Functional").title(),
                tr_priority.get(c.get("priority", "P2"), "Medium"),
                _fmt_preconditions(c.get("preconditions", [])),
                _fmt_steps(c.get("steps", [])),
                c.get("expected", ""),
                f"{c.get('requirement_id', '')} / {c.get('id', '')}",
            ])
    return path


# ──────────────────────────────────────────────────────────────────────────────
# 5. HTML report (self-contained, human-readable)
# ──────────────────────────────────────────────────────────────────────────────
def export_html(cases: list[dict], out_dir: Path) -> Path:
    path = out_dir / "test_cases.html"
    by_prio = Counter(c.get("priority", "?") for c in cases)
    by_type = Counter(c.get("type", "?") for c in cases)
    by_target = Counter(c.get("target", "?") for c in cases)

    rows = []
    for c in cases:
        prio = c.get("priority", "P2")
        prio_class = {"P0": "p0", "P1": "p1", "P2": "p2"}.get(prio, "")
        target_class = f"t-{c.get('target', 'x')}"
        steps_html = "".join(
            f"<li>{html.escape(s.get('action', ''))}"
            + (f" <code>{html.escape(str(s['data']))}</code>" if s.get("data") else "")
            + "</li>"
            for s in c.get("steps", [])
        )
        pcs = c.get("preconditions") or []
        pcs_html = "".join(f"<li>{html.escape(p)}</li>" for p in pcs)

        rows.append(f"""
        <details class="case">
          <summary>
            <span class="id">{html.escape(c.get('id', ''))}</span>
            <span class="badge {prio_class}">{html.escape(prio)}</span>
            <span class="badge {target_class}">{html.escape(c.get('target', ''))}</span>
            <span class="badge t-{html.escape(c.get('type', ''))}">{html.escape(c.get('type', ''))}</span>
            <span class="title">{html.escape(c.get('title', ''))}</span>
            <span class="req">{html.escape(c.get('requirement_id', ''))}</span>
          </summary>
          <div class="body">
            {f'<h4>Preconditions</h4><ul>{pcs_html}</ul>' if pcs_html else ''}
            <h4>Steps</h4>
            <ol>{steps_html}</ol>
            <h4>Expected</h4>
            <p>{html.escape(c.get('expected', ''))}</p>
          </div>
        </details>""")

    page = f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8">
<title>Test Cases ({len(cases)})</title>
<style>
  :root {{ --fg:#1a1a1a; --muted:#666; --bg:#fafafa; --card:#fff; --border:#e0e0e0; }}
  body {{ font-family: -apple-system, 'Segoe UI', system-ui, sans-serif;
          margin:0; padding:2rem; background:var(--bg); color:var(--fg); }}
  h1 {{ margin:0 0 1rem; }}
  .stats {{ display:flex; gap:2rem; flex-wrap:wrap; background:var(--card);
            padding:1rem 1.5rem; border-radius:8px; border:1px solid var(--border); margin-bottom:1.5rem; }}
  .stat-block strong {{ display:block; color:var(--muted); font-size:.8rem; text-transform:uppercase; }}
  .stat-block span {{ margin-right:.75rem; }}
  .case {{ background:var(--card); border:1px solid var(--border); border-radius:6px;
           margin-bottom:.5rem; padding:.5rem .75rem; }}
  .case summary {{ cursor:pointer; display:flex; align-items:center; gap:.5rem; flex-wrap:wrap; }}
  .case summary .id {{ font-family: monospace; font-weight:bold; min-width:70px; }}
  .case summary .title {{ flex:1; min-width:200px; }}
  .case summary .req {{ color:var(--muted); font-size:.85rem; }}
  .badge {{ display:inline-block; padding:.15rem .5rem; border-radius:3px; font-size:.75rem;
            font-weight:600; text-transform:uppercase; letter-spacing:.5px; }}
  .badge.p0 {{ background:#E74C3C; color:#fff; }}
  .badge.p1 {{ background:#F39C12; color:#fff; }}
  .badge.p2 {{ background:#95A5A6; color:#fff; }}
  .badge.t-ui {{ background:#3498DB; color:#fff; }}
  .badge.t-api {{ background:#9B59B6; color:#fff; }}
  .badge.t-manual {{ background:#7F8C8D; color:#fff; }}
  .badge[class*="t-security"] {{ background:#C0392B; color:#fff; }}
  .badge[class*="t-accessibility"] {{ background:#16A085; color:#fff; }}
  .body {{ padding:.5rem 1rem; border-top:1px solid var(--border); margin-top:.5rem; }}
  .body h4 {{ margin:.5rem 0 .25rem; color:var(--muted); font-size:.85rem; text-transform:uppercase; }}
  code {{ background:#f0f0f0; padding:1px 5px; border-radius:3px; font-size:.85em; }}
</style></head><body>
<h1>Test Cases <span style="color:var(--muted); font-weight:normal;">({len(cases)})</span></h1>
<div class="stats">
  <div class="stat-block"><strong>Priority</strong>{"".join(f'<span><b>{p}</b>: {n}</span>' for p, n in by_prio.most_common())}</div>
  <div class="stat-block"><strong>Type</strong>{"".join(f'<span><b>{t}</b>: {n}</span>' for t, n in by_type.most_common())}</div>
  <div class="stat-block"><strong>Target</strong>{"".join(f'<span><b>{t}</b>: {n}</span>' for t, n in by_target.most_common())}</div>
</div>
{"".join(rows)}
</body></html>"""
    path.write_text(page, encoding="utf-8")
    return path


# ──────────────────────────────────────────────────────────────────────────────
# 6. Markdown — summary table + full detail
# ──────────────────────────────────────────────────────────────────────────────
def export_markdown(cases: list[dict], out_dir: Path) -> Path:
    path = out_dir / "test_cases.md"
    lines = [f"# Test Cases ({len(cases)})\n"]

    # Summary stats
    lines.append("## Summary\n")
    lines.append(f"- **Priority:** " + ", ".join(
        f"{p}: {n}" for p, n in Counter(c.get('priority', '?') for c in cases).most_common()))
    lines.append(f"- **Type:** " + ", ".join(
        f"{t}: {n}" for t, n in Counter(c.get('type', '?') for c in cases).most_common()))
    lines.append(f"- **Target:** " + ", ".join(
        f"{t}: {n}" for t, n in Counter(c.get('target', '?') for c in cases).most_common()))
    lines.append("")

    # Overview table
    lines.append("## Overview\n")
    lines.append("| ID | REQ | Title | Type | Target | Priority | Auto |")
    lines.append("|----|-----|-------|------|--------|----------|------|")
    for c in cases:
        title = c.get("title", "").replace("|", "\\|")
        lines.append(
            f"| {c.get('id', '')} | {c.get('requirement_id', '')} | {title} | "
            f"{c.get('type', '')} | {c.get('target', '')} | {c.get('priority', '')} | "
            f"{'✓' if c.get('automatable', True) else '—'} |"
        )
    lines.append("")

    # Full detail
    lines.append("## Details\n")
    for c in cases:
        lines.append(f"### {c.get('id', '')}: {c.get('title', '')}\n")
        lines.append(f"**Requirement:** {c.get('requirement_id', '')}  ")
        lines.append(f"**Type:** {c.get('type', '')}  |  "
                     f"**Target:** {c.get('target', '')}  |  "
                     f"**Priority:** {c.get('priority', '')}")
        pcs = c.get("preconditions") or []
        if pcs:
            lines.append("\n**Preconditions:**")
            lines += [f"- {p}" for p in pcs]
        lines.append("\n**Steps:**")
        for i, s in enumerate(c.get("steps", []), 1):
            extra = f" _(data: `{s['data']}`)_" if s.get("data") else ""
            lines.append(f"{i}. {s.get('action', '')}{extra}")
        lines.append(f"\n**Expected:** {c.get('expected', '')}\n")
        lines.append("---\n")

    path.write_text("\n".join(lines), encoding="utf-8")
    return path


# ──────────────────────────────────────────────────────────────────────────────
# Test Plan exporters — PDF / DOCX / HTML
# ──────────────────────────────────────────────────────────────────────────────
def export_plan(plan_md: str, out_dir: Path, formats: list[str]) -> dict[str, Path]:
    """Export the test plan markdown to PDF/DOCX/HTML. Returns {format: path}."""
    results = {}
    handlers = {
        "html": export_plan_html,
        "pdf": export_plan_pdf,
        "docx": export_plan_docx,
    }
    for fmt in formats:
        fn = handlers.get(fmt)
        if not fn:
            continue
        try:
            results[fmt] = fn(plan_md, out_dir)
        except SystemExit:
            raise
        except Exception as e:  # noqa: BLE001
            print(f"  ⚠ plan {fmt} export failed: {e}")
    return results


def export_plan_html(plan_md: str, out_dir: Path) -> Path:
    """Render the test plan markdown as a self-contained styled HTML page."""
    try:
        import markdown as md_lib
    except ImportError:
        raise SystemExit("HTML plan export needs: pip install markdown")

    body = md_lib.markdown(
        plan_md,
        extensions=["tables", "fenced_code", "toc", "sane_lists"],
    )
    path = out_dir / "test_plan.html"
    page = f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8">
<title>Test Plan</title>
<style>
  :root {{ --fg:#1a1a1a; --muted:#666; --bg:#fafafa; --card:#fff; --accent:#2C3E50;
          --border:#e5e5e5; }}
  body {{ font-family: -apple-system,"Segoe UI",system-ui,sans-serif;
          max-width: 820px; margin: 2rem auto; padding: 0 1.5rem;
          background: var(--bg); color: var(--fg); line-height: 1.6; }}
  h1 {{ color: var(--accent); border-bottom: 2px solid var(--accent); padding-bottom:.4rem;
        margin-top:0; font-size: 1.9rem; }}
  h2 {{ color: var(--accent); margin-top: 2rem; font-size: 1.4rem;
        border-bottom: 1px solid var(--border); padding-bottom: .25rem; }}
  h3 {{ color: var(--accent); font-size: 1.1rem; margin-top: 1.5rem; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1rem 0;
           background: var(--card); border-radius: 6px; overflow: hidden; }}
  th, td {{ padding: .55rem .8rem; text-align: left; border-bottom: 1px solid var(--border); }}
  th {{ background: var(--accent); color: white; font-weight: 600; }}
  tr:last-child td {{ border-bottom: none; }}
  code {{ background: #f0f0f0; padding: 1px 5px; border-radius: 3px; font-size: .9em; }}
  ul, ol {{ padding-left: 1.5rem; }}
  li {{ margin: .25rem 0; }}
  blockquote {{ border-left: 3px solid var(--accent); padding: .25rem 1rem;
                color: var(--muted); margin: 1rem 0; background: #f5f5f5; }}
</style></head><body>{body}</body></html>"""
    path.write_text(page, encoding="utf-8")
    return path


def export_plan_pdf(plan_md: str, out_dir: Path) -> Path:
    """Render the test plan as a PDF using reportlab (pure Python, no OS deps)."""
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
        )
    except ImportError:
        raise SystemExit("PDF plan export needs: pip install reportlab")

    path = out_dir / "test_plan.pdf"
    doc = SimpleDocTemplate(
        str(path), pagesize=LETTER,
        rightMargin=0.75 * inch, leftMargin=0.75 * inch,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    accent = HexColor("#2C3E50")
    styles.add(ParagraphStyle("H1Custom", parent=styles["Heading1"],
                              textColor=accent, spaceAfter=12, fontSize=18))
    styles.add(ParagraphStyle("H2Custom", parent=styles["Heading2"],
                              textColor=accent, spaceAfter=8, fontSize=14))
    styles.add(ParagraphStyle("H3Custom", parent=styles["Heading3"],
                              textColor=accent, spaceAfter=6, fontSize=12))
    styles.add(ParagraphStyle("BodyCustom", parent=styles["BodyText"],
                              fontSize=10.5, leading=14, spaceAfter=6))

    story = []
    for block in _parse_md_blocks(plan_md):
        kind, text = block
        if kind == "h1":
            story.append(Paragraph(_escape_pdf(text), styles["H1Custom"]))
        elif kind == "h2":
            story.append(Paragraph(_escape_pdf(text), styles["H2Custom"]))
        elif kind == "h3":
            story.append(Paragraph(_escape_pdf(text), styles["H3Custom"]))
        elif kind == "table":
            tbl = _md_table_to_reportlab(text)
            if tbl:
                story.append(tbl)
                story.append(Spacer(1, 8))
        elif kind == "list":
            for item in text.split("\n"):
                story.append(Paragraph(f"• {_escape_pdf(item.lstrip('- ').lstrip('* '))}",
                                       styles["BodyCustom"]))
        elif kind == "p" and text.strip():
            story.append(Paragraph(_escape_pdf(text), styles["BodyCustom"]))
        elif kind == "blank":
            story.append(Spacer(1, 6))

    doc.build(story)
    return path


def export_plan_docx(plan_md: str, out_dir: Path) -> Path:
    """Render the test plan as a .docx using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401
    except ImportError:
        raise SystemExit("DOCX plan export needs: pip install python-docx")

    path = out_dir / "test_plan.docx"
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    accent = RGBColor(0x2C, 0x3E, 0x50)

    for kind, text in _parse_md_blocks(plan_md):
        if kind == "h1":
            h = doc.add_heading(_strip_md(text), level=1)
            for run in h.runs:
                run.font.color.rgb = accent
        elif kind == "h2":
            h = doc.add_heading(_strip_md(text), level=2)
            for run in h.runs:
                run.font.color.rgb = accent
        elif kind == "h3":
            h = doc.add_heading(_strip_md(text), level=3)
            for run in h.runs:
                run.font.color.rgb = accent
        elif kind == "table":
            rows = _md_table_rows(text)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
                doc.add_paragraph()
        elif kind == "list":
            for item in text.split("\n"):
                doc.add_paragraph(_strip_md(item.lstrip("- ").lstrip("* ")),
                                  style="List Bullet")
        elif kind == "p" and text.strip():
            doc.add_paragraph(_strip_md(text))

    doc.save(path)
    return path


# ─── Markdown → block parsing helpers (minimal, no full-MD parser needed) ────
def _parse_md_blocks(md: str) -> list[tuple[str, str]]:
    """Convert a markdown string into a list of (kind, text) block tuples.

    A minimal, dependency-free parser — we only need to handle the structures
    our test plan prompts produce (headings h1-h3, tables, bullet lists,
    paragraphs). Does not handle code fences, blockquotes, or nested lists.

    Returns a list of (kind, text) where kind is one of:
      'h1', 'h2', 'h3' — heading levels
      'table'          — full markdown table block (header + separator + rows)
      'list'           — consecutive bullet lines joined with \n
      'p'              — paragraph (consecutive non-special lines joined with space)
      'blank'          — empty line (used as a spacer by PDF/DOCX builders)
    """
    blocks = []
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        if not stripped:
            blocks.append(("blank", ""))
            i += 1
            continue

        if stripped.startswith("# "):
            blocks.append(("h1", stripped[2:].strip()))
            i += 1
        elif stripped.startswith("## "):
            blocks.append(("h2", stripped[3:].strip()))
            i += 1
        elif stripped.startswith("### "):
            blocks.append(("h3", stripped[4:].strip()))
            i += 1
        elif "|" in stripped and i + 1 < len(lines) and "---" in lines[i + 1]:
            # Markdown table: collect header + separator + all following pipe rows
            tbl = [stripped]
            i += 1  # advance to separator row
            tbl.append(lines[i])
            i += 1
            while i < len(lines) and "|" in lines[i]:
                tbl.append(lines[i])
                i += 1
            blocks.append(("table", "\n".join(tbl)))
        elif stripped.startswith(("- ", "* ")):
            # Bullet list: collect consecutive lines starting with - or *
            items = [stripped]
            i += 1
            while i < len(lines) and lines[i].startswith(("- ", "* ")):
                items.append(lines[i])
                i += 1
            blocks.append(("list", "\n".join(items)))
        else:
            # Paragraph: gather consecutive non-blank, non-heading, non-list lines
            para = [stripped]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith(
                ("#", "- ", "* ", "|")
            ):
                para.append(lines[i].strip())
                i += 1
            blocks.append(("p", " ".join(para)))

    return blocks


def _md_table_rows(md_table: str) -> list[list[str]]:
    """Extract rows from a markdown table string."""
    lines = [ln for ln in md_table.splitlines() if "|" in ln and "---" not in ln]
    rows = []
    for ln in lines:
        parts = [p.strip() for p in ln.strip("|").split("|")]
        rows.append([_strip_md(p) for p in parts])
    return rows


def _md_table_to_reportlab(md_table: str):
    """Convert a markdown table to a reportlab Table with accent header."""
    try:
        from reportlab.platypus import Table, TableStyle
        from reportlab.lib.colors import HexColor, whitesmoke
    except ImportError:
        return None

    rows = _md_table_rows(md_table)
    if not rows:
        return None

    # Wrap cells in Paragraphs for wrapping
    from reportlab.platypus import Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    styles = getSampleStyleSheet()
    small = styles["BodyText"]
    wrapped = [[Paragraph(_escape_pdf(cell), small) for cell in row] for row in rows]

    t = Table(wrapped, repeatRows=1, hAlign="LEFT")
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#2C3E50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), whitesmoke),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
        ("TOPPADDING", (0, 0), (-1, 0), 6),
        ("GRID", (0, 0), (-1, -1), 0.25, HexColor("#CCCCCC")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    return t


def _strip_md(s: str) -> str:
    """Strip inline markdown markers (**bold**, *italic*, `code`)."""
    import re as _re
    s = _re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = _re.sub(r"\*(.+?)\*", r"\1", s)
    s = _re.sub(r"`(.+?)`", r"\1", s)
    return s


def _escape_pdf(s: str) -> str:
    """Escape HTML-ish chars for reportlab Paragraph + convert simple MD."""
    import re as _re
    s = _strip_md(s)
    s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return s